"""
docx_export.py — генерация форматированного DOCX для «Цифрового правозащитника».

Оформление по ГОСТ Р 7.0.97-2016 (деловые документы):
  - Шрифт: Times New Roman 12pt
  - Поля: левое 30мм, правое 15мм, верх/низ 20мм
  - Межстрочный интервал: 1.5
  - Отступ первой строки: 1.25 см
  - Выравнивание тела: по ширине
"""

from __future__ import annotations
import io
import re

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Константы стиля
# ─────────────────────────────────────────────
FONT_NAME  = "Times New Roman"
FONT_SIZE  = Pt(12)
LINE_SPACE = Pt(18)   # 1.5 × 12pt

# Заголовки документов (по ним определяем тип и форматирование)
TITLE_KEYWORDS = ("ПРЕТЕНЗИЯ", "ИСКОВОЕ ЗАЯВЛЕНИЕ", "ЗАЯВЛЕНИЕ")

# Строки-разделители, которые идут в шапку (справа)
HEADER_PREFIXES = (
    "кому:", "от:", "адрес:", "истец:", "ответчик:",
    "в ", "г.", "ул.", "пр.", "пер.",
)


# ─────────────────────────────────────────────
# Вспомогательные функции
# ─────────────────────────────────────────────
def _set_run_font(run, bold: bool = False, size: Pt = FONT_SIZE):
    run.font.name  = FONT_NAME
    run.font.size  = size
    run.font.bold  = bold
    # Кириллица требует явного указания шрифта через rFonts
    rpr = run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"),    FONT_NAME)
    rfonts.set(qn("w:hAnsi"),    FONT_NAME)
    rfonts.set(qn("w:cs"),       FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rpr.insert(0, rfonts)


def _set_paragraph_spacing(para, space_before: Pt = Pt(0), space_after: Pt = Pt(6)):
    pf = para.paragraph_format
    pf.space_before    = space_before
    pf.space_after     = space_after
    pf.line_spacing    = LINE_SPACE


def _add_paragraph(doc: Document, text: str,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   bold: bool = False,
                   indent: bool = True,
                   space_before: Pt = Pt(0),
                   space_after: Pt = Pt(6)) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    _set_paragraph_spacing(para, space_before=space_before, space_after=space_after)
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    run = para.add_run(text)
    _set_run_font(run, bold=bold)


def _add_empty_line(doc: Document) -> None:
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, space_before=Pt(0), space_after=Pt(0))
    run = para.add_run("")
    _set_run_font(run)


# ─────────────────────────────────────────────
# Парсер структуры документа
# ─────────────────────────────────────────────
def _classify_line(line: str) -> str:
    """
    Определяет тип строки:
      'title'   — ПРЕТЕНЗИЯ / ИСКОВОЕ ЗАЯВЛЕНИЕ
      'header'  — шапка (кому / от кого / реквизиты)
      'heading' — подзаголовок (ПРОШУ, ТРЕБУЮ, Приложения)
      'list'    — элемент списка (1. / - / •)
      'sign'    — строка подписи (дата + ФИО + Подпись)
      'empty'   — пустая строка
      'body'    — обычный абзац
    """
    stripped = line.strip()

    if not stripped:
        return "empty"

    # Заголовок документа
    if any(stripped.upper().startswith(kw) for kw in TITLE_KEYWORDS):
        return "title"

    # Строка подписи (содержит "Подпись" или серию пробелов + дата)
    if "подпись" in stripped.lower() or re.match(r"^\d{2}[./]\d{2}[./]\d{4}", stripped):
        return "sign"

    # Подзаголовки секций (ПРОШУ, ТРЕБУЮ, Приложения и т.п.)
    if re.match(r"^(ПРОШУ|ТРЕБУЮ|Приложени|ПРИЛОЖЕНИ)", stripped):
        return "heading"

    # Элементы нумерованного / маркированного списка
    if re.match(r"^(\d+[.)]\s|[-•–]\s)", stripped):
        return "list"

    # Шапка документа
    lower = stripped.lower()
    if any(lower.startswith(p) for p in HEADER_PREFIXES):
        return "header"

    # Короткая строка без точки в конце — вероятно тоже шапка
    if len(stripped) < 80 and not stripped.endswith(".") and ":" in stripped:
        return "header"

    return "body"


# ─────────────────────────────────────────────
# Основная функция генерации
# ─────────────────────────────────────────────
def generate_docx(doc_text: str, doc_type: str = "") -> bytes:
    """
    Принимает текст документа (строка), возвращает bytes готового .docx.

    doc_type — «Досудебная претензия» или «Исковое заявление» (для колонтитула).
    """
    document = Document()

    # ── Поля страницы (ГОСТ) ─────────────────
    section = document.sections[0]
    section.page_width  = Mm(210)   # A4
    section.page_height = Mm(297)
    section.left_margin   = Mm(30)
    section.right_margin  = Mm(15)
    section.top_margin    = Mm(20)
    section.bottom_margin = Mm(20)

    # ── Стиль Normal по умолчанию ────────────
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    # ── Нижний колонтитул с номером страницы ─
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run()
    _set_run_font(footer_run, size=Pt(10))
    # Поле PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    footer_run._r.append(fldChar1)
    footer_run._r.append(instrText)
    footer_run._r.append(fldChar2)

    # ── Парсим и рендерим строки ─────────────
    lines = doc_text.split("\n")

    # Накапливаем блок шапки чтобы выровнять его вправо
    header_block: list[str] = []

    def flush_header():
        """Выводим накопленную шапку."""
        if not header_block:
            return
        for h_line in header_block:
            _add_paragraph(
                document, h_line,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
                indent=False,
                space_after=Pt(2),
            )
        header_block.clear()
        _add_empty_line(document)

    i = 0
    while i < len(lines):
        line = lines[i]
        kind = _classify_line(line)

        if kind == "empty":
            flush_header()
            # Пропускаем лишние пустые строки подряд
            if i > 0 and _classify_line(lines[i - 1]) != "empty":
                _add_empty_line(document)

        elif kind == "header":
            header_block.append(line.strip())

        elif kind == "title":
            flush_header()
            _add_empty_line(document)
            _add_paragraph(
                document, line.strip(),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                indent=False,
                space_before=Pt(6),
                space_after=Pt(12),
            )

        elif kind == "heading":
            flush_header()
            _add_empty_line(document)
            _add_paragraph(
                document, line.strip(),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                bold=True,
                indent=False,
                space_before=Pt(6),
                space_after=Pt(4),
            )

        elif kind == "list":
            flush_header()
            _add_paragraph(
                document, line.strip(),
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                indent=False,
                space_after=Pt(3),
            )

        elif kind == "sign":
            flush_header()
            _add_empty_line(document)
            _add_paragraph(
                document, line.strip(),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                indent=False,
                space_after=Pt(3),
            )

        else:  # body
            flush_header()
            _add_paragraph(
                document, line.strip(),
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                indent=True,
                space_after=Pt(6),
            )

        i += 1

    flush_header()

    # ── Сохраняем в bytes ────────────────────
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
